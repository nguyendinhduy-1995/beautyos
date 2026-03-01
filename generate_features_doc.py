#!/usr/bin/env python3
"""Generate BeautyOS 29 PRO Features DOCX from JSON data"""
import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TIER = {1:'Ưu tiên cao',2:'Phát triển',3:'Sắp ra mắt',4:'AI-Powered'}
TIER_COUNT = {1:0,2:0,3:0,4:0}

def shade(cell, color):
    s = OxmlElement('w:shd')
    s.set(qn('w:fill'), color)
    s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)

def add_feature(doc, idx, f):
    h = doc.add_heading(level=2)
    r = h.add_run(f"{idx}. {f['name']}")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x0f,0x17,0x2a)

    t = doc.add_table(rows=4, cols=2)
    t.style = 'Table Grid'
    info = [('ID',f['id']),('Phân loại',TIER.get(f['tier'],'')),
            ('Giá',f"{f['price']:,}đ/tháng".replace(',','.')),
            ('Phổ biến','✅ Có' if f.get('popular') else '—')]
    for i,(l,v) in enumerate(info):
        c1,c2 = t.cell(i,0), t.cell(i,1)
        c1.text, c2.text = l, v
        c1.paragraphs[0].runs[0].font.bold = True
        c1.paragraphs[0].runs[0].font.size = Pt(10)
        c2.paragraphs[0].runs[0].font.size = Pt(10)
        shade(c1, 'F1F5F9')
        c1.width, c2.width = Cm(3.5), Cm(12)
    doc.add_paragraph('')

    # Mô tả
    p = doc.add_paragraph()
    p.add_run('Mô tả:').bold = True
    p = doc.add_paragraph(f['desc'])
    for r in p.runs: r.font.size = Pt(10)

    # Tính năng
    p = doc.add_paragraph()
    p.add_run('Tính năng chi tiết:').bold = True
    for feat in f['features']:
        p = doc.add_paragraph(feat, style='List Bullet')
        for r in p.runs: r.font.size = Pt(10)

    # Giao diện
    p = doc.add_paragraph()
    p.add_run('Giao diện (UI/UX):').bold = True
    p = doc.add_paragraph(f['ui'])
    for r in p.runs: r.font.size = Pt(10)

    # Quy trình hoạt động
    p = doc.add_paragraph()
    p.add_run('Quy trình hoạt động:').bold = True
    for step in f['workflow']:
        p = doc.add_paragraph(step, style='List Bullet')
        for r in p.runs: r.font.size = Pt(10)

    # Separator
    p = doc.add_paragraph()
    r = p.add_run('─' * 60)
    r.font.color.rgb = RGBColor(0xe2,0xe8,0xf0)
    r.font.size = Pt(8)

def main():
    with open('/Volumes/Data - 3/beautyos/features_data.json','r') as fp:
        features = json.load(fp)

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    # Title
    title = doc.add_heading(level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run('BeautyOS — 29 Tính Năng PRO')
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0x0f,0x17,0x2a)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('Cấu trúc tính năng · Giao diện · Quy trình hoạt động')
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x64,0x74,0x8b)
    doc.add_paragraph('')

    # Overview
    doc.add_heading('Tổng quan', level=1)
    doc.add_paragraph('BeautyOS Premium: 29 module add-on riêng lẻ, 4 nhóm. '
        'Kích hoạt riêng từng module, dùng thử 14 ngày miễn phí, hủy bất cứ lúc nào. Giá từ 150.000đ/tháng.')

    # Count tiers
    for f in features: TIER_COUNT[f['tier']] += 1

    st = doc.add_table(rows=5, cols=4)
    st.style = 'Table Grid'
    for i,h in enumerate(['Nhóm','Số lượng','Giá từ','Ghi chú']):
        c = st.cell(0,i)
        c.text = h
        c.paragraphs[0].runs[0].font.bold = True
        c.paragraphs[0].runs[0].font.size = Pt(10)
        shade(c, '1E293B')
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xff,0xff,0xff)

    rows_data = [
        ('Tier 1 — Ưu tiên cao','4','190.000đ','AI Tư vấn, SMS/ZNS, Báo cáo, Loyalty'),
        ('Tier 2 — Phát triển','4','150.000đ','CRM Automation, Đa chi nhánh, KPI, Bảo mật'),
        ('Tier 3 — Sắp ra mắt','8','190.000đ','Before/After, EMR, Booking Online, CTV, Kho, App, Pixel'),
        ('Tier 4 — AI-Powered','13','190.000đ','Ads Meta AI, Content AI, Phân tích Da AI, Telesales AI...'),
    ]
    for i,(a,b,c,d) in enumerate(rows_data):
        for j,v in enumerate([a,b,c,d]):
            cell = st.cell(i+1,j)
            cell.text = v
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_page_break()

    # Group and render by tier
    by_tier = {1:[],2:[],3:[],4:[]}
    for f in features: by_tier[f['tier']].append(f)

    idx = 1
    for tn in [1,2,3,4]:
        tier_list = by_tier[tn]
        h = doc.add_heading(level=1)
        h.add_run(f'TIER {tn} — {TIER[tn].upper()}').font.size = Pt(16)
        doc.add_paragraph(f'{len(tier_list)} tính năng trong nhóm này.')
        doc.add_paragraph('')
        for f in tier_list:
            add_feature(doc, idx, f)
            idx += 1
        if tn < 4: doc.add_page_break()

    out = '/Volumes/Data - 3/beautyos/BeautyOS_29_Pro_Features.docx'
    doc.save(out)
    print(f'✅ Saved: {out}')

if __name__ == '__main__':
    main()
